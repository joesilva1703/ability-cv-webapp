#!/usr/bin/env python3
"""
fill_cv.py — Populate the Ability Group Master CV Template with candidate data.

Usage:
    python fill_cv.py <candidate_data.json> <output.docx> [--template master_template.docx]

The JSON schema is documented in SKILL.md. Missing fields are rendered as
"TBC" with a yellow highlight so the recruiter can spot them at a glance.
"""
from __future__ import annotations

import argparse
import copy
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

TBC = "TBC"
BODY_FONT = "Century Gothic"
BODY_SIZE_PT = 9

# Section-banner / "main heading" text. These paragraphs keep their original
# (inherited) size; only the font name is normalised.
MAIN_HEADING_TEXTS = {
    "CANDIDATE SUMMARY",
    "PERSONAL INFORMATION",
    "WORK SUMMARY",
    "EDUCATION",
    "TERTIARY EDUCATION",
    "PROFESSIONAL MEMBERSHIPS",
    "SKILLS",
    "COMPUTER SKILLS",
    "WORK EXPERIENCE",
}


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _copy_run_formatting(src_run, dst_run) -> None:
    """Copy character-level formatting from src_run to dst_run."""
    src_rpr = src_run._element.find(qn("w:rPr"))
    if src_rpr is not None:
        dst_rpr = dst_run._element.find(qn("w:rPr"))
        if dst_rpr is not None:
            dst_run._element.remove(dst_rpr)
        dst_run._element.insert(0, copy.deepcopy(src_rpr))


def _clear_paragraph(paragraph) -> None:
    """Remove all runs (text) from a paragraph but keep its properties."""
    for child in list(paragraph._element):
        if child.tag == qn("w:pPr"):
            continue
        paragraph._element.remove(child)


def _set_paragraph_text(paragraph, text: str, *, tbc: bool = False,
                        bold: bool | None = None) -> None:
    """Replace all text in a paragraph while preserving formatting of first run."""
    # Preserve first run's formatting if possible
    template_run = paragraph.runs[0] if paragraph.runs else None
    _clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    if template_run is not None:
        _copy_run_formatting(template_run, run)
    if bold is not None:
        run.bold = bold
    if tbc:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _write_value(cell, value: str | None, *, bold: bool = False) -> None:
    """Write a value into a cell, using TBC highlight if value is missing."""
    is_missing = value is None or (isinstance(value, str) and not value.strip())
    text = TBC if is_missing else str(value).strip()
    # Use the first paragraph; clear any extra paragraphs after
    if not cell.paragraphs:
        cell.add_paragraph()
    first_p = cell.paragraphs[0]
    _set_paragraph_text(first_p, text, tbc=is_missing, bold=bold)
    # Remove subsequent paragraphs that existed in the template
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)


def _clone_element_after(element):
    """Deep-copy an XML element and insert it immediately after the original."""
    new = copy.deepcopy(element)
    element.addnext(new)
    return new


def _add_bullet_paragraphs(cell, items: list[str], *, template_idx: int = 0,
                           tbc_if_empty: bool = False) -> None:
    """Replace a cell's contents with a bulleted list using the cell's existing
    bullet paragraph as the style template. Blank items are dropped, and an
    empty list leaves the cell empty (no TBC placeholder) unless the caller
    explicitly opts in via ``tbc_if_empty=True``."""
    # Drop any blank / whitespace-only bullets so we never render an empty bullet.
    items = [str(s) for s in (items or []) if s is not None and str(s).strip()]
    if not items:
        if tbc_if_empty:
            items = [TBC]
        else:
            # Clear the cell so no stray empty bullet shows
            for p in list(cell.paragraphs):
                p._element.getparent().remove(p._element)
            cell.add_paragraph()  # keep one empty paragraph so the cell renders
            return
    # Find a paragraph in the cell to use as a style template (bulleted style)
    paragraphs = list(cell.paragraphs)
    template_p = paragraphs[template_idx] if paragraphs else None
    template_el = template_p._element if template_p is not None else None
    # Remove every paragraph except the template (compare by XML element identity)
    for p in paragraphs:
        if p._element is template_el:
            continue
        p._element.getparent().remove(p._element)
    if template_p is None:
        # Fall back: plain paragraphs
        for item in items:
            p = cell.add_paragraph()
            _set_paragraph_text(p, item, tbc=(item == TBC))
        return
    # First item goes into template_p; clone template for subsequent items
    _set_paragraph_text(template_p, items[0], tbc=(items[0] == TBC))
    last_el = template_el
    from docx.text.paragraph import Paragraph
    for item in items[1:]:
        new_el = _clone_element_after(last_el)
        new_p = Paragraph(new_el, template_p._parent)
        _set_paragraph_text(new_p, item, tbc=(item == TBC))
        last_el = new_el


# ---------------------------------------------------------------------------
# Section fillers — each targets a specific table by index (from analysis)
# ---------------------------------------------------------------------------

def fill_header(doc, candidate: dict) -> None:
    """Table 0: logo cell containing 'Candidate Name' and 'Position' paragraphs."""
    cell = doc.tables[0].rows[0].cells[0]
    # The cell has: logo paragraph, candidate name paragraph, position paragraph.
    name_p = None
    pos_p = None
    for p in cell.paragraphs:
        text = p.text.strip()
        if text == "Candidate Name":
            name_p = p
        elif text == "Position" and name_p is not None:
            pos_p = p
            break
    if name_p is not None:
        _set_paragraph_text(name_p, candidate.get("name") or TBC,
                            tbc=not candidate.get("name"))
    if pos_p is not None:
        _set_paragraph_text(pos_p, candidate.get("position") or TBC,
                            tbc=not candidate.get("position"))


def fill_introducer(doc, introducer: dict) -> None:
    """Paragraph 0: 'INTRODUCED BY: Consultant name | 010 593 4900 | email address'."""
    for p in doc.paragraphs:
        if p.text.strip().startswith("INTRODUCED BY"):
            name = introducer.get("name") or TBC
            phone = introducer.get("phone") or TBC
            email = introducer.get("email") or TBC
            # Preserve the first run's formatting (bold "INTRODUCED BY:")
            template_run = p.runs[0] if p.runs else None
            _clear_paragraph(p)
            prefix = p.add_run("INTRODUCED BY: ")
            if template_run is not None:
                _copy_run_formatting(template_run, prefix)
            prefix.bold = True
            body = p.add_run(f"{name} | {phone} | {email}")
            body.bold = True
            return


def add_top_disclaimer_spacer(doc) -> None:
    """Insert a blank paragraph after the top 'Disclaimer:' block (the
    centred title with the embedded textbox) so there is a visible gap
    before the CANDIDATE SUMMARY heading. The last-page disclaimer's
    heading is plain 'Disclaimer' (no colon) and is not affected."""
    for p in doc.paragraphs:
        if p.text.strip() == "Disclaimer:":
            spacer = OxmlElement("w:p")
            p._element.addnext(spacer)
            return


def fill_summary(doc, summary: dict | None) -> None:
    """Table 2: CANDIDATE SUMMARY content cell. Renders only the bullet
    points — the narrative ``paragraph`` field is intentionally ignored.
    A blank paragraph is inserted above the first bullet so there is a
    visible gap between the CANDIDATE SUMMARY heading and the bullets."""
    cell = doc.tables[2].rows[0].cells[0]
    summary = summary or {}
    bullets = summary.get("bullets") or []
    _add_bullet_paragraphs(cell, bullets)
    # Spacer paragraph at the top of the cell.
    spacer = OxmlElement("w:p")
    if cell.paragraphs:
        cell.paragraphs[0]._element.addprevious(spacer)
    else:
        cell._tc.append(spacer)
        cell.add_paragraph()


def fill_personal_info(doc, candidate: dict) -> None:
    """Table 3: 6 rows x 4 cols. Data cells are (row, col): (2,1)(2,3)(3,1)(3,3)(4,1)(4,3)(5,1)."""
    tbl = doc.tables[3]
    mapping = {
        (2, 1): candidate.get("location"),
        (2, 3): candidate.get("ee_gender"),
        (3, 1): candidate.get("availability"),
        (3, 3): candidate.get("nationality"),
        (4, 1): candidate.get("current_salary"),
        (4, 3): candidate.get("id_number"),
        (5, 1): candidate.get("expected_salary"),
    }
    for (r, c), value in mapping.items():
        _write_value(tbl.rows[r].cells[c], value)


def fill_work_summary(doc, jobs: list[dict]) -> None:
    """Table 5: 2 rows x 3 cols. Row 1 has 3 paragraphs per cell (company / position / dates)."""
    tbl = doc.tables[5]
    row = tbl.rows[1]
    columns = ["company", "position", "dates"]
    if not jobs:
        jobs = [{c: None for c in columns}]
    from docx.text.paragraph import Paragraph
    for ci, key in enumerate(columns):
        cell = row.cells[ci]
        paragraphs = list(cell.paragraphs)
        template_p = paragraphs[0] if paragraphs else None
        template_el = template_p._element if template_p is not None else None
        # Remove every paragraph except the template (compare by XML element)
        for p in paragraphs:
            if p._element is template_el:
                continue
            p._element.getparent().remove(p._element)
        if template_p is None:
            continue
        first_val = jobs[0].get(key)
        _set_paragraph_text(template_p, first_val or TBC, tbc=not first_val)
        last_el = template_el
        for job in jobs[1:]:
            val = job.get(key)
            new_el = _clone_element_after(last_el)
            new_p = Paragraph(new_el, template_p._parent)
            _set_paragraph_text(new_p, val or TBC, tbc=not val)
            last_el = new_el


def fill_school(doc, school: dict) -> None:
    """Table 7: School Name (date) + Matric line."""
    tbl = doc.tables[7]
    cell = tbl.rows[0].cells[0]
    # Cell has 2 paragraphs: "School Name (date)" and "Matric"
    paragraphs = list(cell.paragraphs)
    if len(paragraphs) >= 1:
        name = school.get("name") if school else None
        date = school.get("date") if school else None
        if name and date:
            line = f"{name} ({date})"
        elif name:
            line = name
        else:
            line = TBC
        _set_paragraph_text(paragraphs[0], line, tbc=not name, bold=True)
    if len(paragraphs) >= 2:
        matric = (school or {}).get("matric") or TBC
        _set_paragraph_text(paragraphs[1], matric, tbc=not (school or {}).get("matric"))


def fill_tertiary(doc, tertiary: list[dict]) -> None:
    """Table 8: 4 rows. Row 2 has Institution (date) + qualification. Clone for each."""
    tbl = doc.tables[8]
    # Row indices: 0=TERTIARY header, 1=blank, 2=Institution+qualification, 3=PROFESSIONAL MEMBERSHIPS
    template_row = tbl.rows[2]
    if not tertiary:
        # Render single row with TBC
        cell = template_row.cells[0]
        paragraphs = list(cell.paragraphs)
        if len(paragraphs) >= 1:
            _set_paragraph_text(paragraphs[0], TBC, tbc=True, bold=True)
        if len(paragraphs) >= 2:
            _set_paragraph_text(paragraphs[1], TBC, tbc=True, bold=True)
        return
    # Fill first entry into template_row
    def _fill_tert_row(row, entry):
        cell = row.cells[0]
        paragraphs = list(cell.paragraphs)
        inst = entry.get("institution") or TBC
        date = entry.get("date")
        line1 = f"{inst} ({date})" if date and entry.get("institution") else inst
        qual = entry.get("qualification") or TBC
        if len(paragraphs) >= 1:
            _set_paragraph_text(paragraphs[0], line1,
                                tbc=not entry.get("institution"))
        if len(paragraphs) >= 2:
            _set_paragraph_text(paragraphs[1], qual,
                                tbc=not entry.get("qualification"), bold=True)
    _fill_tert_row(template_row, tertiary[0])
    # Clone for each additional entry, inserting before the PROFESSIONAL MEMBERSHIPS row
    last_row_el = template_row._tr
    for entry in tertiary[1:]:
        new_row_el = copy.deepcopy(template_row._tr)
        last_row_el.addnext(new_row_el)
        # Wrap for python-docx
        from docx.table import _Row
        new_row = _Row(new_row_el, tbl)
        _fill_tert_row(new_row, entry)
        last_row_el = new_row_el


def fill_memberships(doc, memberships: list[str]) -> None:
    """Table 9: single cell holding bulleted memberships."""
    tbl = doc.tables[9]
    cell = tbl.rows[0].cells[0]
    _add_bullet_paragraphs(cell, memberships or [], tbc_if_empty=False)
    # If empty and the cell has no paragraph, add one blank
    if not memberships and not cell.paragraphs:
        cell.add_paragraph()


def fill_skills(doc, skills: list[str]) -> None:
    """Table 11: single cell, bulleted skills list."""
    tbl = doc.tables[11]
    cell = tbl.rows[0].cells[0]
    _add_bullet_paragraphs(cell, skills or [])


def fill_computer_skills(doc, computer_skills: list[str]) -> None:
    """Table 12: row 2 holds bulleted computer skills."""
    tbl = doc.tables[12]
    cell = tbl.rows[2].cells[0]
    _add_bullet_paragraphs(cell, computer_skills or [])


def fill_work_experience(doc, experiences: list[dict]) -> None:
    """Tables 14+ : each job uses a 6-row table. Template has 2 sample tables."""
    # Gather the two template tables
    job_tables = [doc.tables[14], doc.tables[15]]
    if not experiences:
        experiences = [{}]
    # If we have fewer experiences than template tables, remove extras
    while len(job_tables) > len(experiences):
        extra = job_tables.pop()
        extra._element.getparent().remove(extra._element)
    # If more experiences than templates, clone the last template for each additional.
    # Insert a blank paragraph between consecutive role tables so there's a
    # visible gap between roles (the template already has one between TBL14
    # and TBL15; new tables need their own spacer added).
    while len(job_tables) < len(experiences):
        src = job_tables[-1]
        spacer_el = OxmlElement("w:p")
        src._element.addnext(spacer_el)
        new_tbl_el = copy.deepcopy(src._element)
        spacer_el.addnext(new_tbl_el)
        from docx.table import Table
        new_tbl = Table(new_tbl_el, src._parent)
        job_tables.append(new_tbl)

    for tbl, exp in zip(job_tables, experiences):
        # Row 0: company (underlined bold)
        _write_value(tbl.rows[0].cells[0], exp.get("company"), bold=True)
        # Row 1: position
        _write_value(tbl.rows[1].cells[0], exp.get("position"), bold=True)
        # Row 2: dates
        _write_value(tbl.rows[2].cells[0], exp.get("dates"))
        # Row 3: blank (separator) — leave as-is
        # Row 4: Duties (bulleted list)
        duties_cell = tbl.rows[4].cells[0]
        _add_bullet_paragraphs(duties_cell, exp.get("duties") or [])
        # Row 5: 'Reason for leaving: <reason>' — preserve the bold label, append reason
        reason_cell = tbl.rows[5].cells[0]
        reason = exp.get("reason_for_leaving")
        _fill_reason(reason_cell, reason)


def _fill_reason(cell, reason: str | None) -> None:
    """Row 5 in a work-experience table. Preserve 'Reason for leaving:' bold label
    (even when split across multiple runs) and append the reason text (or TBC)."""
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.add_paragraph()
        paragraphs = list(cell.paragraphs)
    p = paragraphs[0]
    first_el = p._element
    # Accumulate text from all runs up to and including the colon — that's the label
    full_text = "".join(r.text for r in p.runs)
    if ":" in full_text:
        label_text = full_text.split(":", 1)[0].strip() + ":"
    else:
        label_text = full_text.strip() or "Reason for leaving:"
    template_run = p.runs[0] if p.runs else None
    _clear_paragraph(p)
    label_run = p.add_run(label_text + " ")
    if template_run is not None:
        _copy_run_formatting(template_run, label_run)
    label_run.bold = True
    is_missing = not reason
    reason_run = p.add_run(reason if reason else TBC)
    if is_missing:
        reason_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # Remove extra paragraphs (anything after the first one in the cell)
    for extra_p in list(cell.paragraphs):
        if extra_p._element is first_el:
            continue
        extra_p._element.getparent().remove(extra_p._element)


# ---------------------------------------------------------------------------
# Body-font normalisation
# ---------------------------------------------------------------------------

def _set_run_font(run, *, name: str, size_pt: int | None) -> None:
    """Force a run's font name (all script slots) and optionally its size."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    # Set every script slot so themed fonts (asciiTheme/hAnsiTheme) are overridden.
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    # Remove any themed-font attributes that would otherwise win over the explicit name.
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _is_protected_paragraph(p) -> bool:
    """Return True if this paragraph should be left completely untouched.
    Matches the two disclaimers by text. Handles the top disclaimer's
    'Disclaimer:' centered title (with colon) and the last-page disclaimer's
    'Disclaimer' heading + body."""
    text = p.text.strip()
    if not text:
        return False
    if text in ("Disclaimer", "Disclaimer:"):
        return True
    if text.startswith("You may only use this information"):
        return True
    if text.startswith("The information contained in this document"):
        return True
    return False


def _apply_body_font(doc) -> None:
    """Normalise body content to Century Gothic 9pt. The following are left
    completely untouched (font and size unchanged):
      * candidate name and position (TBL0 row 0 col 0 — the title block)
      * the 'Disclaimer:' title at the top of the CV (and its embedded textbox)
      * the disclaimer on the last page ('Disclaimer' heading + body text)
    Main section headings (CANDIDATE SUMMARY, PERSONAL INFORMATION, WORK
    SUMMARY, EDUCATION, TERTIARY EDUCATION, PROFESSIONAL MEMBERSHIPS, SKILLS,
    COMPUTER SKILLS, WORK EXPERIENCE) keep their current size; only the font
    name is normalised to Century Gothic."""
    def _process_paragraph(p):
        if _is_protected_paragraph(p):
            return
        is_main_heading = p.text.strip().upper() in MAIN_HEADING_TEXTS
        for run in p.runs:
            _set_run_font(
                run,
                name=BODY_FONT,
                size_pt=None if is_main_heading else BODY_SIZE_PT,
            )

    def _walk_cell(cell):
        for p in cell.paragraphs:
            _process_paragraph(p)
        for nested in cell.tables:
            _walk_table(nested, skip_first_cell=False)

    def _walk_table(tbl, *, skip_first_cell: bool):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                if skip_first_cell and ri == 0 and ci == 0:
                    # Candidate name + position cell — leave entirely untouched.
                    continue
                _walk_cell(cell)

    for p in doc.paragraphs:
        _process_paragraph(p)

    for ti, tbl in enumerate(doc.tables):
        # Only the very first table's first cell is the title block; nested
        # tables (if any) elsewhere should be processed normally.
        _walk_table(tbl, skip_first_cell=(ti == 0))

    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is None:
                continue
            for p in part.paragraphs:
                _process_paragraph(p)
            for tbl in part.tables:
                _walk_table(tbl, skip_first_cell=False)


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------

def fill_cv(template_path: Path, data: dict, output_path: Path) -> None:
    doc = Document(str(template_path))
    fill_header(doc, data.get("candidate") or {})
    fill_introducer(doc, data.get("introducer") or {})
    add_top_disclaimer_spacer(doc)
    fill_summary(doc, data.get("summary"))
    fill_personal_info(doc, data.get("candidate") or {})
    fill_work_summary(doc, data.get("work_summary") or [])
    fill_school(doc, data.get("school") or {})
    fill_tertiary(doc, data.get("tertiary") or [])
    fill_memberships(doc, data.get("memberships") or [])
    fill_skills(doc, data.get("skills") or [])
    fill_computer_skills(doc, data.get("computer_skills") or [])
    fill_work_experience(doc, data.get("work_experience") or [])
    _apply_body_font(doc)
    doc.save(str(output_path))


def main():
    ap = argparse.ArgumentParser(description="Fill the Ability Group master CV template.")
    ap.add_argument("data_json", help="Path to JSON file with candidate data.")
    ap.add_argument("output_docx", help="Path where the filled .docx will be written.")
    ap.add_argument("--template", default=str(Path(__file__).parent / "master_template.docx"),
                    help="Path to the master template (default: bundled master_template.docx).")
    args = ap.parse_args()
    data = json.loads(Path(args.data_json).read_text())
    fill_cv(Path(args.template), data, Path(args.output_docx))
    print(f"Wrote {args.output_docx}")


if __name__ == "__main__":
    main()
