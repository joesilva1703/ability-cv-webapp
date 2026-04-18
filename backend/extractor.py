"""
Extract plain text from an uploaded CV (PDF or .docx).

Tries multiple backends in order:

PDF:
  1. pypdf (pure-Python, bundled)
  2. pdftotext with -layout (only if installed on the machine)
  3. pandoc (only if installed)

.docx:
  1. python-docx paragraphs + tables
  2. pandoc fallback

Returns a single string. Raises EmptyExtractionError if nothing was extractable
(typically means the PDF is an image-only scan).
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path


class EmptyExtractionError(RuntimeError):
    """Raised when no text could be pulled from the CV."""


def extract_text(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".pdf"}:
        text = _extract_pdf(file_bytes)
    elif suffix in {".docx", ".doc"}:
        text = _extract_docx(file_bytes, suffix)
    else:
        raise ValueError(
            f"Unsupported file type: {suffix!r}. Upload a PDF or Word document."
        )
    text = _normalise(text)
    if not text.strip():
        raise EmptyExtractionError(
            "No text could be read from the CV. If it's a scanned image, "
            "please send a text-selectable version."
        )
    return text


def _normalise(text: str) -> str:
    # Collapse runs of >2 blank lines; strip trailing whitespace per line.
    lines = [ln.rstrip() for ln in text.replace("\r", "").split("\n")]
    out = []
    blank = 0
    for ln in lines:
        if ln.strip() == "":
            blank += 1
            if blank <= 2:
                out.append("")
        else:
            blank = 0
            out.append(ln)
    return "\n".join(out).strip()


def _extract_pdf(data: bytes) -> str:
    # 1. pypdf
    try:
        from pypdf import PdfReader  # type: ignore

        import io

        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
        text = "\n\n".join(pages)
        if text.strip():
            return text
    except Exception:
        pass

    # 2. pdftotext -layout
    if shutil.which("pdftotext"):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "cv.pdf"
            src.write_bytes(data)
            try:
                out = subprocess.run(
                    ["pdftotext", "-layout", str(src), "-"],
                    check=True,
                    capture_output=True,
                )
                text = out.stdout.decode("utf-8", errors="ignore")
                if text.strip():
                    return text
            except Exception:
                pass

    # 3. pandoc
    if shutil.which("pandoc"):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "cv.pdf"
            src.write_bytes(data)
            dst = Path(td) / "cv.md"
            try:
                subprocess.run(
                    ["pandoc", str(src), "-o", str(dst)],
                    check=True,
                    capture_output=True,
                )
                text = dst.read_text(encoding="utf-8", errors="ignore")
                if text.strip():
                    return text
            except Exception:
                pass

    return ""


def _extract_docx(data: bytes, suffix: str) -> str:
    # 1. python-docx
    if suffix == ".docx":
        try:
            import io

            from docx import Document  # type: ignore

            doc = Document(io.BytesIO(data))
            chunks = []
            for p in doc.paragraphs:
                if p.text.strip():
                    chunks.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells]
                    line = " | ".join([c for c in cells if c])
                    if line:
                        chunks.append(line)
            text = "\n".join(chunks)
            if text.strip():
                return text
        except Exception:
            pass

    # 2. pandoc fallback (handles .doc too if present)
    if shutil.which("pandoc"):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / f"cv{suffix}"
            src.write_bytes(data)
            dst = Path(td) / "cv.md"
            try:
                subprocess.run(
                    ["pandoc", str(src), "-o", str(dst)],
                    check=True,
                    capture_output=True,
                )
                text = dst.read_text(encoding="utf-8", errors="ignore")
                if text.strip():
                    return text
            except Exception:
                pass

    return ""
